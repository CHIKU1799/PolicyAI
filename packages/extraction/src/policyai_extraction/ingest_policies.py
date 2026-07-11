"""Ingest a firm's OWN policy documents into its knowledge base (CompanyDocument).

This is the company-side counterpart to ``ingest.py`` (which ingests regulations).
The obligation mapper reviews every applicable regulation against these documents
(``mapping._nearest_kb_docs``), so loading the firm's real policies here is what
makes the gap analysis concrete rather than blind.

Give it files or directories of the firm's policies (PDF / DOCX / TXT / MD):

    make ingest-policies DIR=inbox/policies
    uv run python -m policyai_extraction.ingest_policies inbox/policies [--dry-run]
    uv run python -m policyai_extraction.ingest_policies --backfill   # re-embed only

For each file: extract text, embed (bge-m3 / Cohere), and upsert a CompanyDocument.
Dedup is on content_hash, so re-running is idempotent. Scanned PDFs with no text
layer are flagged ``needs_ocr`` rather than embedded empty (which would poison
similarity search). ``--backfill`` re-embeds any existing doc whose embedding is
NULL — use it to repair rows inserted outside this path.
"""

from __future__ import annotations

import argparse
import asyncio
import hashlib
import io
import logging
import sys
from pathlib import Path
from uuid import UUID

from policyai_graph.db import make_engine, make_sessionmaker
from policyai_graph.models_app import DEFAULT_ORG_ID, CompanyDocument, DocumentStatus
from sqlalchemy import select
from sqlalchemy.ext.asyncio import AsyncSession

from policyai_extraction.embeddings import embed_text

log = logging.getLogger("policyai.ingest_policies")

# File types we can pull text from (see policyai_api.textextract).
SUFFIXES = {".pdf", ".docx", ".txt", ".md"}
# Below this many characters we assume no real text layer (scanned image PDF).
MIN_TEXT = 50
# The mapper embeds/searches on the leading slice; keep it bounded like the API route.
EMBED_CHARS = 8000


def _extract_text(content: bytes, *, filename: str) -> str:
    """Plain text from a PDF / DOCX / text file. Mirrors the API's textextract but
    kept local so this package doesn't depend on policyai_api. Scanned PDFs with no
    text layer yield ~nothing, so the caller flags them needs_ocr instead of
    embedding empty."""
    name = filename.lower()
    if name.endswith(".pdf"):
        from pypdf import PdfReader

        reader = PdfReader(io.BytesIO(content))
        return "\n".join((p.extract_text() or "") for p in reader.pages).strip()
    if name.endswith(".docx"):
        from docx import Document

        doc = Document(io.BytesIO(content))
        return "\n".join(p.text for p in doc.paragraphs).strip()
    return content.decode("utf-8", errors="ignore").strip()


def _mime_for(name: str) -> str | None:
    n = name.lower()
    if n.endswith(".pdf"):
        return "application/pdf"
    if n.endswith(".docx"):
        return "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    if n.endswith(".txt"):
        return "text/plain"
    if n.endswith(".md"):
        return "text/markdown"
    return None


async def _embed_with_retry(text: str, *, tries: int = 3) -> list[float]:
    """Embed, retrying transient provider/network hiccups with linear backoff."""
    last: Exception | None = None
    for i in range(tries):
        try:
            return await embed_text(text[:EMBED_CHARS])
        except Exception as exc:  # noqa: BLE001 - provider/network transient
            last = exc
            await asyncio.sleep(1.5 * (i + 1))
    assert last is not None
    raise last


def _iter_files(paths: list[Path]) -> list[Path]:
    out: list[Path] = []
    for p in paths:
        if p.is_dir():
            out.extend(sorted(f for f in p.rglob("*") if f.suffix.lower() in SUFFIXES))
        elif p.is_file() and p.suffix.lower() in SUFFIXES:
            out.append(p)
        else:
            log.warning("skipping (not a supported file/dir): %s", p)
    return out


async def ingest_paths(
    session: AsyncSession,
    files: list[Path],
    *,
    org_id: UUID = DEFAULT_ORG_ID,
    dry_run: bool = False,
) -> dict:
    """Load each policy file into CompanyDocument. Commits per file so a mid-batch
    failure keeps everything already done. Dedup on content_hash."""
    tally = {"embedded": 0, "needs_ocr": 0, "skipped": 0, "failed": 0}
    for f in files:
        try:
            content = f.read_bytes()
        except Exception as exc:  # noqa: BLE001
            tally["failed"] += 1
            log.warning("read failed %s: %s", f.name, exc)
            continue
        text = _extract_text(content, filename=f.name)
        content_hash = hashlib.sha256(content).hexdigest()

        existing = (
            await session.execute(
                select(CompanyDocument).where(
                    CompanyDocument.org_id == org_id,
                    CompanyDocument.content_hash == content_hash,
                )
            )
        ).scalar_one_or_none()
        # Already loaded AND already embedded → nothing to do.
        if existing is not None and existing.embedding is not None:
            tally["skipped"] += 1
            continue
        if dry_run:
            tally["embedded" if len(text) >= MIN_TEXT else "needs_ocr"] += 1
            continue

        doc = existing or CompanyDocument(org_id=org_id)
        doc.storage_path = f"local:{f}"
        doc.filename = f.name
        doc.mime = _mime_for(f.name)
        doc.content_hash = content_hash
        doc.raw_text = text
        if existing is None:
            session.add(doc)

        if len(text) < MIN_TEXT:
            doc.status = DocumentStatus.NEEDS_OCR.value
            tally["needs_ocr"] += 1
            log.warning("no text layer (needs OCR): %s", f.name)
        else:
            try:
                doc.embedding = await _embed_with_retry(text)
                doc.status = DocumentStatus.PROCESSED.value
                tally["embedded"] += 1
                log.info("embedded policy: %s (%d chars)", f.name, len(text))
            except Exception as exc:  # noqa: BLE001
                doc.status = DocumentStatus.FAILED.value
                tally["failed"] += 1
                log.warning("embed failed %s: %s", f.name, exc)
        await session.commit()
    return tally


async def backfill_embeddings(session: AsyncSession, *, org_id: UUID = DEFAULT_ORG_ID) -> dict:
    """Re-embed every CompanyDocument that has text but a NULL embedding — repairs
    rows inserted outside the normal embed path (the cause of blind gap analysis)."""
    tally = {"embedded": 0, "needs_ocr": 0, "failed": 0}
    docs = (
        (
            await session.execute(
                select(CompanyDocument).where(
                    CompanyDocument.org_id == org_id,
                    CompanyDocument.embedding.is_(None),
                )
            )
        )
        .scalars()
        .all()
    )
    log.info("backfill: %d company docs with NULL embedding", len(docs))
    for doc in docs:
        text = (doc.raw_text or "").strip()
        if len(text) < MIN_TEXT:
            doc.status = DocumentStatus.NEEDS_OCR.value
            tally["needs_ocr"] += 1
            continue
        try:
            doc.embedding = await _embed_with_retry(text)
            doc.status = DocumentStatus.PROCESSED.value
            tally["embedded"] += 1
            log.info("backfilled: %s", doc.filename)
        except Exception as exc:  # noqa: BLE001
            doc.status = DocumentStatus.FAILED.value
            tally["failed"] += 1
            log.warning("backfill embed failed %s: %s", doc.filename, exc)
        await session.commit()
    return tally


async def _run(paths: list[Path], *, org_id: UUID, dry_run: bool, backfill: bool) -> dict:
    engine = make_engine()
    sm = make_sessionmaker(engine)
    try:
        async with sm() as session:
            if backfill:
                return await backfill_embeddings(session, org_id=org_id)
            files = _iter_files(paths)
            log.info("%d policy files to process", len(files))
            return await ingest_paths(session, files, org_id=org_id, dry_run=dry_run)
    finally:
        await engine.dispose()


def main() -> int:
    ap = argparse.ArgumentParser(description="Ingest a firm's policy docs into its KB.")
    ap.add_argument("paths", type=Path, nargs="*", help="policy files or directories")
    ap.add_argument("--org-id", type=str, default=None, help="target org UUID (default demo org)")
    ap.add_argument("--dry-run", action="store_true", help="report counts without writing")
    ap.add_argument(
        "--backfill",
        action="store_true",
        help="re-embed existing company docs with NULL embedding (ignores paths)",
    )
    args = ap.parse_args()
    logging.basicConfig(level=logging.INFO, format="%(levelname)s %(name)s: %(message)s")
    if not args.paths and not args.backfill:
        ap.error("provide at least one path, or --backfill")
    org_id = UUID(args.org_id) if args.org_id else DEFAULT_ORG_ID
    tally = asyncio.run(
        _run(args.paths, org_id=org_id, dry_run=args.dry_run, backfill=args.backfill)
    )
    print(tally)
    return 0 if tally.get("failed", 0) == 0 else 1


if __name__ == "__main__":
    sys.exit(main())
